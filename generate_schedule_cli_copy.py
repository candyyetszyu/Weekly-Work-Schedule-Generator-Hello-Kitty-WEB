import pandas as pd
import random
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches

def get_monday(date_str):
    date = datetime.strptime(date_str, "%Y-%m-%d")
    return date - timedelta(days=date.weekday())

def split_weekly_minutes(total_minutes):
    """Safely split total weekly minutes across 7 days with 30–120 min/day using 30-min units."""
    min_day = 30
    max_day = 120
    max_weekly = 15 * 60
    unit = 30  # 30-minute units

    if total_minutes > max_weekly:
        raise ValueError("Total weekly minutes must not exceed 900 (15 hours).")
    if total_minutes < min_day:
        raise ValueError("Total weekly minutes must be at least 30.")

    # Round total minutes to nearest 30-minute unit
    total_minutes = round(total_minutes / unit) * unit

    # Step 1: Start with 0 for all days
    result = [0] * 7
    day_indices = list(range(7))
    random.shuffle(day_indices)

    # Step 2: Pre-fill with min 30 mins per day until used or out of budget
    remaining = total_minutes
    for i in day_indices:
        if remaining >= min_day:
            result[i] = min_day
            remaining -= min_day
        else:
            break

    # Step 3: Distribute remaining in 30-min blocks, capped at 120
    attempts = 0
    max_attempts = 100  # Prevent infinite loop
    
    while remaining >= unit and attempts < max_attempts:
        attempts += 1
        random.shuffle(day_indices)
        distributed_this_round = False
        
        for i in day_indices:
            if result[i] + unit <= max_day:
                result[i] += unit
                remaining -= unit
                distributed_this_round = True
                if remaining < unit:
                    break
        
        # If we couldn't distribute any minutes this round, break to avoid infinite loop
        if not distributed_this_round:
            break
    
    # If we still have remaining minutes, distribute them as best we can
    if remaining > 0:
        for i in range(7):
            if result[i] + remaining <= max_day:
                result[i] += remaining
                remaining = 0
                break

    return result

def random_start_time(duration_min):
    """Generate a random time between 09:00 and 18:00 that fits the session, aligned to :00 or :30."""
    earliest_start = 9 * 60
    latest_start = 18 * 60 - duration_min

    if latest_start < earliest_start:
        raise ValueError(f"Duration {duration_min} min too long to fit within working hours (09:00–18:00).")

    valid_slots = [t for t in range(earliest_start, latest_start + 1, 30)]
    if not valid_slots:
        raise ValueError(f"No valid slots found for {duration_min} minutes between 09:00 and 18:00.")

    start_min = random.choice(valid_slots)
    hour = start_min // 60
    minute = start_min % 60
    return f"{hour:02d}:{minute:02d}"

def generate_schedule(any_date_str, total_hours_per_week, total_days, output_filename, start_week, 
                     export_txt=True, export_xlsx=True, export_docx=True):
    start_date = get_monday(any_date_str)
    total_minutes_per_week = round(total_hours_per_week * 60 / 30) * 30  # Round to 30-minute units
    weekdays = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]

    text_lines = []
    excel_data = []
    docx_data = []  # For .docx table format
    current_date = start_date
    end_date = start_date + timedelta(days=total_days - 1)
    week_num = start_week - 1
    total_minutes_accumulated = 0

    while current_date <= end_date:
        week_num += 1
        week_start = current_date
        week_end = min(current_date + timedelta(days=6), end_date)
        text_lines.append(f"Week {week_num}: {week_start.strftime('%d %B')} – {week_end.strftime('%d %B')}")

        # Generate daily work minutes
        week_minutes = split_weekly_minutes(total_minutes_per_week)
        
        # Collect weekly schedule details for .docx table
        week_schedule_details = []

        # Adjust to fit remaining days if not a full week
        days_in_this_week = min(7, (end_date - current_date).days + 1)
        current_sum = sum(week_minutes[:days_in_this_week])
        difference = total_minutes_per_week - current_sum
        week_minutes[days_in_this_week - 1] += difference
        total_minutes_accumulated += sum(week_minutes[:days_in_this_week])

        for i in range(days_in_this_week):
            minutes = week_minutes[i]
            if minutes == 0:
                current_date += timedelta(days=1)
                continue

            hours = minutes // 60
            mins = minutes % 60
            try:
                start_time = random_start_time(minutes)
                end_time = (datetime.strptime(start_time, "%H:%M") + timedelta(minutes=minutes)).strftime("%H:%M")
            except ValueError:
                start_time = "09:00"
                end_time = (datetime.strptime(start_time, "%H:%M") + timedelta(minutes=minutes)).strftime("%H:%M")

            date_str = current_date.strftime("%d %b %Y")
            weekday = weekdays[current_date.weekday()]
            time_str = f"{hours}h {mins}m"
            text_lines.append(f"{date_str} ({weekday}) - {time_str} | {start_time}–{end_time}")
            excel_data.append({
                "Week": f"Week {week_num}",
                "Date": date_str,
                "Day": weekday,
                "Work Time": time_str,
                "Time Slot": f"{start_time}–{end_time}"
            })
            
            # Collect for .docx table
            week_schedule_details.append(f"{date_str} ({weekday}): {time_str} | {start_time}–{end_time}")
            current_date += timedelta(days=1)

        text_lines.append(f"Total hours this week: {total_hours_per_week:.2f}h\n")
        
        # Add weekly data to .docx table - all days in one row per week
        # Extract date and time parts for each column
        date_parts = []
        schedule_times = []
        hours_parts = []
        
        for day_detail in week_schedule_details:
            # Format: "15 Jan 2024 (Monday): 1h 25m | 13:00–14:25"
            parts = day_detail.split(": ")
            date_part = parts[0]  # "15 Jan 2024 (Monday)"
            time_part = parts[1].split(" | ")[1]  # "13:00–14:25"
            hours_part = parts[1].split(" | ")[0]  # "1h 25m"
            
            date_parts.append(date_part)
            schedule_times.append(time_part)
            hours_parts.append(hours_part)
        
        docx_data.append({
            "Week": f"Week {week_num}",
            "Date": "\n".join(date_parts),
            "Schedule": "\n".join(schedule_times),
            "Hours": "\n".join(hours_parts)
        })

    total_hours_final = total_minutes_accumulated / 60
    text_lines.append("==============================")
    text_lines.append(f"Total work time: {total_hours_final:.2f} hours")
    text_lines.append("==============================")

    excel_data.append({
        "Week": "",
        "Date": "",
        "Day": "Total work time (hours)",
        "Work Time": f"{total_hours_final:.2f}",
        "Time Slot": ""
    })

    # Output
    created_files = []
    
    if export_txt:
        txt_file = f"{output_filename}.txt"
        with open(txt_file, "w") as f:
            f.write("\n".join(text_lines))
        created_files.append(txt_file)
    
    if export_xlsx:
        xlsx_file = f"{output_filename}.xlsx"
        pd.DataFrame(excel_data).to_excel(xlsx_file, index=False)
        created_files.append(xlsx_file)
    
    if export_docx:
        docx_file = f"{output_filename}.docx"
        # Create .docx table
        doc = Document()
        doc.add_heading('Weekly Work Schedule', 0)
        
        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Week'
        header_cells[1].text = 'Date'
        header_cells[2].text = 'Schedule'
        header_cells[3].text = 'Hours'
        
        # Add data rows
        for week_data in docx_data:
            row_cells = table.add_row().cells
            row_cells[0].text = week_data['Week']
            row_cells[1].text = week_data['Date']
            row_cells[2].text = week_data['Schedule']
            row_cells[3].text = week_data['Hours']
        
        # Add total row
        total_row = table.add_row().cells
        total_row[0].text = ''
        total_row[1].text = ''
        total_row[2].text = 'Total work time (hours)'
        total_row[3].text = f"{total_hours_final:.2f}"
        
        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(1.0)  # Week
            row.cells[1].width = Inches(2.5)  # Date (more space for individual days)
            row.cells[2].width = Inches(2.5)  # Schedule (time ranges)
            row.cells[3].width = Inches(1.0)  # Hours
        
        doc.save(docx_file)
        created_files.append(docx_file)

    print(f"\n✅ Schedule created from {start_date.strftime('%Y-%m-%d')} for {total_days} days.")
    for file in created_files:
        if file.endswith('.txt'):
            print(f"📄 Text saved to: {file}")
        elif file.endswith('.xlsx'):
            print(f"📊 Excel saved to: {file}")
        elif file.endswith('.docx'):
            print(f"📝 Word document saved to: {file}")
    print(f"🕒 Total work time: {total_hours_final:.2f} hours\n")

    if export_txt and created_files:
        txt_file = next((f for f in created_files if f.endswith('.txt')), None)
        if txt_file:
            with open(txt_file, "r") as f:
                print(f.read())

def generate_schedule_total_hours(any_date_str, total_overall_hours, output_filename, start_week,
                                 export_txt=True, export_xlsx=True, export_docx=True):
    max_weekly_hours = 15
    min_weekly_hours = 0.5  # Minimum 30 minutes per week
    
    if total_overall_hours < min_weekly_hours:
        raise ValueError(f"Total overall hours must be at least {min_weekly_hours} hours.")
    
    # Calculate total minutes and round to nearest 30
    total_minutes = int(round(total_overall_hours * 60 / 30) * 30)
    
    # Calculate how many full weeks we need
    weeks_required = (total_minutes + max_weekly_hours * 60 - 1) // (max_weekly_hours * 60)
    
    # Calculate hours per week (distribute evenly)
    hours_per_week = total_overall_hours / weeks_required
    
    # Ensure we don't exceed max weekly hours
    if hours_per_week > max_weekly_hours:
        hours_per_week = max_weekly_hours
    
    # Calculate total days needed
    total_days = weeks_required * 7
    
    print(f"📊 Distribution: {total_overall_hours:.2f} hours over {weeks_required} weeks")
    print(f"⏰ {hours_per_week:.2f} hours per week")
    print(f"📅 {total_days} days total\n")
    
    generate_schedule(any_date_str, hours_per_week, total_days, output_filename, start_week,
                     export_txt, export_xlsx, export_docx)

if __name__ == "__main__":
    print("🗓 Weekly Work Schedule Generator")
    print("Choose a mode:")
    print("1. Input total **weekly** hours")
    print("2. Input total **overall** hours (system will split into weeks)")
    mode = input("Enter 1 or 2: ").strip()

    any_date = input("Enter any date within the starting week (YYYY-MM-DD): ").strip()
    start_week = int(input("Enter the starting week number (e.g. 1): ").strip())
    filename = input("Enter output filename (no extension): ").strip()

    if mode == "1":
        total_hours = float(input("Enter total weekly work hours (≤ 15): ").strip())
        if total_hours > 15:
            raise ValueError("Weekly hours cannot exceed 15.")
        total_days = int(input("Enter total number of days to generate (e.g. 60): ").strip())
        generate_schedule(any_date, total_hours, total_days, filename, start_week)

    elif mode == "2":
        total_overall_hours = float(input("Enter total overall work hours (e.g. 50): ").strip())
        if total_overall_hours <= 0:
            raise ValueError("Total overall hours must be greater than 0.")
        generate_schedule_total_hours(any_date, total_overall_hours, filename, start_week)

    else:
        print("❌ Invalid mode selected.")